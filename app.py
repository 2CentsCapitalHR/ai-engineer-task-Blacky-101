import streamlit as st
from io import BytesIO
from docx import Document
import json
import re
from datetime import datetime

COMPANY_INCORP_CHECKLIST = [
    "Articles of Association",
    "Memorandum of Association",
    "Incorporation Application Form",
    "UBO Declaration Form",
    "Register of Members and Directors"
]

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Incorporation Application Form": ["incorporation application", "application form"],
    "UBO Declaration Form": ["ubo", "ultimate beneficial owner", "ubo declaration"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"]
}

def parse_docx_bytes(file_bytes):
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return doc, paragraphs

def detect_doc_type(paragraphs):
    text = " ".join(paragraphs).lower()
    found = []
    for dtype, kws in DOC_TYPE_KEYWORDS.items():
        for kw in kws:
            if kw in text:
                found.append(dtype)
                break
    if not found:
        return "Unknown"
    return ", ".join(found)

def checklist_verify(detected_types, process="company_incorporation"):
    required = COMPANY_INCORP_CHECKLIST
    uploaded = detected_types
    missing = [d for d in required if d not in uploaded]
    return {"process": "Company Incorporation",
            "documents_uploaded": len(uploaded),
            "required_documents": len(required),
            "missing_documents": missing}

def detect_red_flags(paragraphs):
    issues = []
    ambiguous_phrases = ["best efforts", "commercially reasonable", "reasonable endeavours", "best endeavours"]
    for i, p in enumerate(paragraphs):
        text = p.lower()
        if re.search(r"\b(u[ae]\b|united arab emirates|federal court)\b", text) and "adgm" not in text:
            issues.append({
                "document_section": f"para_{i}",
                "issue": "Potential incorrect jurisdiction reference (mentions UAE federal courts or ambiguous jurisdiction)",
                "severity": "High",
                "suggestion": "Specify ADGM Courts as jurisdiction where required.",
                "para_index": i
            })
        if re.search(r"(signature|signed by|__________________)", text) is None and i >= len(paragraphs)-3:
            issues.append({
                "document_section": f"para_{i}",
                "issue": "Possible missing signature block near end of document",
                "severity": "Medium",
                "suggestion": "Ensure signatory name, designation, and signature block present.",
                "para_index": i
            })
        for phrase in ambiguous_phrases:
            if phrase in text:
                issues.append({
                    "document_section": f"para_{i}",
                    "issue": f"Ambiguous phrase '{phrase}' found",
                    "severity": "Low",
                    "suggestion": f"Consider replacing '{phrase}' with a specific, binding obligation.",
                    "para_index": i
                })
    seen = set()
    dedup = []
    for it in issues:
        key = (it['para_index'], it['issue'])
        if key not in seen:
            dedup.append(it); seen.add(key)
    return dedup

def add_comments_to_docx(doc: Document, issues):
    for it in issues:
        idx = it.get("para_index", None)
        if idx is None: continue
        if idx < 0 or idx >= len(doc.paragraphs): continue
        paragraph = doc.paragraphs[idx]
        try:
            paragraph.add_run(f"\n[REVIEW SUGGESTION] {it['suggestion']}")
        except Exception as e:
            pass
    return doc

st.title("Corporate Agent — Prototype (ADGM compliance helper)")

uploaded_files = st.file_uploader("Upload one or more .docx files (company formation docs recommended)", accept_multiple_files=True, type=['docx'])

if uploaded_files:
    st.write(f"Uploaded {len(uploaded_files)} file(s). Processing...")
    processed = []
    detected_types = []
    all_issues = []

    for uploaded in uploaded_files:
        name = uploaded.name
        bytes_data = uploaded.read()
        doc, paragraphs = parse_docx_bytes(bytes_data)

        doc_type = detect_doc_type(paragraphs)
        detected_types.append(doc_type)
        issues = detect_red_flags(paragraphs)
        reviewed_doc = add_comments_to_docx(doc, issues)

        out_stream = BytesIO()
        reviewed_doc.save(out_stream)
        out_stream.seek(0)

        st.markdown(f"**{name}** — detected type(s): {doc_type}")
        st.download_button(label=f"Download reviewed: {name}",
                           data=out_stream.getvalue(),
                           file_name=f"reviewed_{name}",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        processed.append({
            "filename": name,
            "detected_type": doc_type,
            "issues_found": issues
        })
        all_issues.extend([{"file": name, **it} for it in issues])

    checklist_result = checklist_verify([t for t in detected_types if t != "Unknown"])
    st.subheader("Checklist verification (company incorporation)")
    st.json(checklist_result)

    summary = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "process": checklist_result["process"],
        "documents_uploaded": checklist_result["documents_uploaded"],
        "required_documents": checklist_result["required_documents"],
        "missing_documents": checklist_result["missing_documents"],
        "issues_found": all_issues
    }
    st.subheader("Structured report")
    st.json(summary)

    st.download_button("Download JSON report", data=json.dumps(summary, indent=2), file_name="report.json", mime="application/json")
else:
    st.info("Upload .docx files to begin. (Prototype adds paragraph-level review suggestions.)")
